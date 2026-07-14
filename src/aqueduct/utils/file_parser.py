"""多格式需求文档解析器。

支持 Markdown / 纯文本 / Excel / PDF / Word 格式的需求文档，
统一转换为 Markdown 文本供后续 LLM 处理。

解析库为延迟导入（lazy import），未安装时仅对应格式不可用，
不影响其他格式的正常工作。
"""

from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_requirement_file(file_path: str | Path) -> str:
    """解析多格式需求文档，返回 Markdown 文本。

    支持的格式：
      - .md / .txt / .markdown → 直接读取文本
      - .xlsx / .xls → 用 openpyxl/pandas 提取表格
      - .pdf → 用 pdfplumber 提取文本和表格
      - .docx → 用 python-docx 提取段落和表格

    Args:
        file_path: 需求文档路径。

    Returns:
        解析后的 Markdown 格式文本。

    Raises:
        FileNotFoundError: 文件不存在。
        ValueError: 不支持的文件格式。
        ImportError: 对应格式的解析库未安装。
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"需求文档不存在: {path}")

    ext = path.suffix.lower()

    if ext in {".md", ".txt", ".markdown"}:
        return path.read_text(encoding="utf-8")

    if ext in {".xlsx", ".xls"}:
        return _parse_excel(path)

    if ext == ".pdf":
        return _parse_pdf(path)

    if ext == ".docx":
        return _parse_docx(path)

    raise ValueError(f"不支持的需求文档格式: {ext}。支持: .md, .txt, .xlsx, .xls, .pdf, .docx")


def _parse_excel(path: Path) -> str:
    """解析 Excel 文件，将每个 sheet 转为 Markdown 表格。"""
    try:
        import openpyxl
    except ImportError:
        raise ImportError("解析 Excel 需要 openpyxl: pip install 'aqueduct[files]'") from None

    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    sections: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        sections.append(f"## {sheet_name}\n")

        # 表头
        header = rows[0]
        header_cells = [str(c) if c is not None else "" for c in header]
        sections.append("| " + " | ".join(header_cells) + " |")
        sections.append("| " + " | ".join(["---"] * len(header_cells)) + " |")

        # 数据行
        for row in rows[1:]:
            cells = [str(c) if c is not None else "" for c in row]
            sections.append("| " + " | ".join(cells) + " |")

        sections.append("")

    wb.close()
    return "\n".join(sections)


def _parse_pdf(path: Path) -> str:
    """解析 PDF 文件，提取文本和表格。"""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("解析 PDF 需要 pdfplumber: pip install 'aqueduct[files]'") from None

    sections: list[str] = []

    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text()
            if text:
                sections.append(f"<!-- Page {i} -->\n")
                sections.append(text)

            tables = page.extract_tables()
            for table in tables:
                if not table:
                    continue
                header = table[0]
                header_cells = [str(c) if c else "" for c in header]
                sections.append("\n| " + " | ".join(header_cells) + " |")
                sections.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
                for row in table[1:]:
                    cells = [str(c) if c else "" for c in row]
                    sections.append("| " + " | ".join(cells) + " |")
            sections.append("")

    return "\n".join(sections)


def _parse_docx(path: Path) -> str:
    """解析 Word 文档，提取段落和表格。"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("解析 Word 需要 python-docx: pip install 'aqueduct[files]'") from None

    doc = Document(str(path))
    sections: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 根据段落样式映射为 Markdown 标题
        style_name = (para.style.name or "").lower()
        if "heading" in style_name:
            try:
                level = int("".join(filter(str.isdigit, style_name)) or "1")
            except ValueError:
                level = 1
            sections.append(f"{'#' * level} {text}")
        else:
            sections.append(text)

    # 提取表格
    for table in doc.tables:
        sections.append("")
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            continue

        header = rows[0]
        sections.append("| " + " | ".join(header) + " |")
        sections.append("| " + " | ".join(["---"] * len(header)) + " |")
        for row in rows[1:]:
            sections.append("| " + " | ".join(row) + " |")
        sections.append("")

    return "\n".join(sections)
